from __future__ import annotations

import json
import re
from io import BytesIO

from docx import Document
from openai import OpenAI
from pypdf import PdfReader

from .config import config
from .models import ResearchProfile


PROFILE_SCHEMA = {
    "type": "object",
    "properties": {
        "name": {"type": "string"},
        "headline": {"type": "string"},
        "domains": {"type": "array", "items": {"type": "string"}},
        "methods": {"type": "array", "items": {"type": "string"}},
        "systems": {"type": "array", "items": {"type": "string"}},
        "current_questions": {"type": "array", "items": {"type": "string"}},
        "adjacent_fields": {"type": "array", "items": {"type": "string"}},
        "keywords": {"type": "array", "items": {"type": "string"}},
    },
    "required": [
        "name",
        "headline",
        "domains",
        "methods",
        "systems",
        "current_questions",
        "adjacent_fields",
        "keywords",
    ],
    "additionalProperties": False,
}


def extract_cv_text(filename: str, payload: bytes) -> str:
    suffix = filename.lower().rsplit(".", 1)[-1]
    if suffix == "pdf":
        reader = PdfReader(BytesIO(payload))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif suffix == "docx":
        document = Document(BytesIO(payload))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                text += "\n" + " | ".join(cell.text for cell in row.cells)
    else:
        raise ValueError("PaperPulse supports PDF and DOCX CV files.")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) < 80:
        raise ValueError("The CV did not contain enough extractable text.")
    return text


def _fallback_profile(text: str) -> ResearchProfile:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    name = next(
        (
            line
            for line in lines[:8]
            if 2 <= len(line.split()) <= 5
            and not any(char.isdigit() for char in line)
            and "@" not in line
        ),
        "Researcher",
    )
    candidates = re.findall(r"\b[A-Za-z][A-Za-z-]{4,}\b", text.lower())
    stop = {
        "research",
        "university",
        "experience",
        "education",
        "publication",
        "publications",
        "department",
        "science",
        "project",
        "projects",
    }
    frequency: dict[str, int] = {}
    for word in candidates:
        if word not in stop:
            frequency[word] = frequency.get(word, 0) + 1
    keywords = [
        word.title()
        for word, _ in sorted(frequency.items(), key=lambda pair: pair[1], reverse=True)[:12]
    ]
    return ResearchProfile(
        name=name,
        headline="Research profile extracted locally — review before ranking",
        domains=keywords[:4],
        methods=keywords[4:8],
        systems=keywords[8:10],
        current_questions=["Track meaningful developments across the listed research areas"],
        adjacent_fields=["Computational methods", "Advanced instrumentation"],
        keywords=keywords,
    )


def build_research_profile(text: str) -> ResearchProfile:
    if not config.openai_api_key:
        return _fallback_profile(text)

    client = OpenAI(api_key=config.openai_api_key)
    response = client.responses.create(
        model=config.analysis_model,
        reasoning={"effort": "low"},
        store=False,
        input=[
            {
                "role": "system",
                "content": (
                    "You create a durable research-interest profile from a CV for a "
                    "scientific literature recommender. Infer current interests from the "
                    "whole trajectory, with extra weight on recent positions, projects, "
                    "papers, methods, and recurring systems. Use concise English noun "
                    "phrases. Do not invent achievements. Make adjacent_fields genuinely "
                    "useful for cross-domain inspiration."
                ),
            },
            {
                "role": "user",
                "content": "Build the research profile from this CV:\n\n" + text[:60_000],
            },
        ],
        text={
            "verbosity": "low",
            "format": {
                "type": "json_schema",
                "name": "research_profile",
                "strict": True,
                "schema": PROFILE_SCHEMA,
            },
        },
    )
    return ResearchProfile.model_validate(json.loads(response.output_text))

